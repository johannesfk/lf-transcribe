from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document

from .asr import ASRResult, Segment
from .prose import Paragraph

logger = logging.getLogger(__name__)


def export_json(path: str | Path, asr: ASRResult) -> None:
    obj = {
        "language": asr.language,
        "segments": [
            {
                "start": s.start,
                "end": s.end,
                "text": s.text,
                "words": s.words,
            }
            for s in asr.segments
        ],
    }
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(json.dumps(obj, indent=2))


def export_markdown(path: str | Path, paragraphs: Iterable[Paragraph]) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    lines: List[str] = []
    for p in paragraphs:
        lines.append(p.text)
        lines.append("")
    Path(path).write_text("\n".join(lines).strip() + "\n")


def export_docx(path: str | Path, paragraphs: Iterable[Paragraph], title: Optional[str] = None) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p.text)
    doc.save(str(path))
